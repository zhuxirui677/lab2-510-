#!/usr/bin/env python3
"""Convert a Markdown file to a .docx using python-docx (no Pandoc required).

Handles common patterns in lab reports: ATX headings, fenced code blocks,
pipe tables, horizontal rules, bullet and numbered lists, blockquotes, and
plain paragraphs. Intended for TECHIN 510 Lab 2 submission bundles.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path


def _try_pandoc(md_path: Path, docx_path: Path) -> bool:
    """Return True if Pandoc successfully wrote docx_path."""
    try:
        subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return True
    except (FileNotFoundError, subprocess.CalledProcessError):
        return False


def _strip_md_inline(text: str) -> str:
    """Remove simple inline Markdown formatting."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def _is_table_sep(line: str) -> bool:
    stripped = line.strip().replace(" ", "")
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    core = stripped.strip("|")
    return bool(core) and all(set(cell) <= {"-", ":"} for cell in core.split("|"))


def convert_markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []

    def add_code_block() -> None:
        if not code_lines:
            return
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        paragraph.paragraph_format.left_indent = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block()
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped == "---" or stripped == "***":
            doc.add_paragraph("—" * 24)
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = _strip_md_inline(heading_match.group(2).strip())
            doc.add_heading(text, level=level)
            idx += 1
            continue

        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote_lines.append(lines[idx].strip().lstrip(">").strip())
                idx += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(_strip_md_inline(" ".join(quote_lines)))
            run.italic = True
            continue

        bullet_match = re.match(r"^[\t ]*([-*])\s+(.*)$", line)
        if bullet_match:
            doc.add_paragraph(_strip_md_inline(bullet_match.group(2)), style="List Bullet")
            idx += 1
            continue

        num_match = re.match(r"^[\t ]*(\d+)\.\s+(.*)$", line)
        if num_match:
            doc.add_paragraph(_strip_md_inline(num_match.group(2)), style="List Number")
            idx += 1
            continue

        if _is_table_row(line) and idx + 1 < len(lines) and _is_table_sep(lines[idx + 1]):
            header = _parse_table_row(line)
            idx += 2
            rows: list[list[str]] = [header]
            while idx < len(lines) and _is_table_row(lines[idx]):
                rows.append(_parse_table_row(lines[idx]))
                idx += 1
            table = doc.add_table(rows=len(rows), cols=len(header))
            table.style = "Table Grid"
            for r, row_cells in enumerate(rows):
                for c, val in enumerate(row_cells):
                    if c < len(table.rows[r].cells):
                        table.rows[r].cells[c].text = _strip_md_inline(val)
            doc.add_paragraph()
            continue

        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if img_match:
            alt, path_str = img_match.group(1), img_match.group(2)
            img_path = (md_path.parent / path_str).resolve()
            if img_path.is_file():
                doc.add_paragraph(alt or "Figure")
                doc.add_picture(str(img_path), width=Inches(6))
            else:
                doc.add_paragraph(f"[Figure omitted — file not found: {path_str}]")
            idx += 1
            continue

        para_lines: list[str] = [stripped]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx]
            ns = nxt.strip()
            if not ns:
                idx += 1
                break
            if (
                ns.startswith("```")
                or ns == "---"
                or ns.startswith("#")
                or ns.startswith(">")
                or re.match(r"^[\t ]*([-*])\s+", nxt)
                or re.match(r"^[\t ]*\d+\.\s+", nxt)
                or _is_table_row(nxt)
                or re.match(r"^!\[", ns)
            ):
                break
            para_lines.append(ns)
            idx += 1
        doc.add_paragraph(_strip_md_inline(" ".join(para_lines)))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX for lab submission.")
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("LAB2_FULL_REPORT.md"),
        help="Markdown source file (default: LAB2_FULL_REPORT.md)",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("LAB2_Submission.docx"),
        help="Output DOCX path (default: LAB2_Submission.docx)",
    )
    args = parser.parse_args()

    root = Path(__file__).resolve().parent.parent
    md_path = args.input if args.input.is_absolute() else root / args.input
    docx_path = args.output if args.output.is_absolute() else root / args.output

    if not md_path.is_file():
        print(f"Input not found: {md_path}", file=sys.stderr)
        return 1

    if _try_pandoc(md_path, docx_path):
        print(f"Wrote {docx_path} (via Pandoc).")
        return 0

    try:
        convert_markdown_to_docx(md_path, docx_path)
    except ImportError:
        print(
            "Missing dependency: install with\n"
            "  pip install -r requirements-docx.txt\n"
            "Or install Pandoc (https://pandoc.org/) for higher-fidelity conversion.",
            file=sys.stderr,
        )
        return 1

    print(f"Wrote {docx_path} (via python-docx).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
